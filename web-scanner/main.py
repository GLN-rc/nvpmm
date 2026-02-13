"""
Website Competitor Scanner - Main FastAPI Application
Analyzes your website, competitors, and brand documents to provide
SEO/GEO/LLM discoverability optimization suggestions.
"""

import os
import io
import tempfile
from typing import Optional
from datetime import datetime
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import uvicorn

from scraper import WebsiteScraper
from analyzer import OptimizationAnalyzer
from document_processor import DocumentProcessor
from metric_explanations import generate_metric_insights, get_all_explanations

app = FastAPI(
    title="Website Competitor Scanner",
    description="Scan your website and competitors to get AI-powered optimization suggestions",
    version="1.0.0"
)

# CORS middleware for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="static"), name="static")


class ScanRequest(BaseModel):
    your_website: str
    competitor_urls: list[str]
    focus_areas: Optional[list[str]] = None


class AnalysisResponse(BaseModel):
    status: str
    your_site_analysis: dict
    competitor_analyses: list[dict]
    recommendations: list[dict]
    priority_actions: list[dict]


@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    """Serve the main frontend application."""
    return FileResponse("static/index.html")


@app.post("/api/scan", response_model=AnalysisResponse)
async def scan_websites(
    your_website: str = Form(...),
    competitor_urls: str = Form(...),  # Comma-separated
    focus_areas: Optional[str] = Form(None),
    brand_docs: list[UploadFile] = File(default=[])
):
    """
    Scan your website and competitor websites, analyze uploaded documents,
    and return prioritized optimization suggestions.
    """
    try:
        scraper = WebsiteScraper()
        analyzer = OptimizationAnalyzer()
        doc_processor = DocumentProcessor()

        # Parse competitor URLs
        competitors = [url.strip() for url in competitor_urls.split(",") if url.strip()]

        # Parse focus areas if provided
        areas = []
        if focus_areas:
            areas = [area.strip() for area in focus_areas.split(",") if area.strip()]

        # Scrape your website
        your_site_data = await scraper.analyze_website(your_website)

        # Scrape competitor websites
        competitor_data = []
        for comp_url in competitors[:5]:  # Limit to 5 competitors
            try:
                comp_analysis = await scraper.analyze_website(comp_url)
                competitor_data.append(comp_analysis)
            except Exception as e:
                competitor_data.append({
                    "url": comp_url,
                    "error": str(e),
                    "status": "failed"
                })

        # Process uploaded documents (temporary, no storage)
        brand_context = []
        for doc in brand_docs:
            if doc.filename:
                # Create temp file, process, then delete
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(doc.filename)[1]) as tmp:
                    content = await doc.read()
                    tmp.write(content)
                    tmp_path = tmp.name

                try:
                    doc_content = doc_processor.extract_content(tmp_path, doc.filename)
                    brand_context.append({
                        "filename": doc.filename,
                        "content": doc_content
                    })
                finally:
                    # Clean up temp file
                    os.unlink(tmp_path)

        # Generate optimization recommendations
        recommendations = await analyzer.generate_recommendations(
            your_site=your_site_data,
            competitors=competitor_data,
            brand_documents=brand_context,
            focus_areas=areas
        )

        # Generate metric insights with explanations
        metric_insights = generate_metric_insights(your_site_data, competitor_data)

        return {
            "status": "success",
            "your_site_analysis": your_site_data,
            "competitor_analyses": competitor_data,
            "recommendations": recommendations["recommendations"],
            "copy_suggestions": recommendations.get("copy_suggestions", []),
            "priority_actions": recommendations["priority_actions"],
            "metric_insights": metric_insights,
            "metric_explanations": get_all_explanations()
        }

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/quick-scan")
async def quick_scan(request: ScanRequest):
    """Quick scan without file uploads - JSON API."""
    scraper = WebsiteScraper()
    analyzer = OptimizationAnalyzer()

    your_site_data = await scraper.analyze_website(request.your_website)

    competitor_data = []
    for comp_url in request.competitor_urls[:5]:
        try:
            comp_analysis = await scraper.analyze_website(comp_url)
            competitor_data.append(comp_analysis)
        except Exception as e:
            competitor_data.append({
                "url": comp_url,
                "error": str(e),
                "status": "failed"
            })

    recommendations = await analyzer.generate_recommendations(
        your_site=your_site_data,
        competitors=competitor_data,
        brand_documents=[],
        focus_areas=request.focus_areas or []
    )

    return {
        "status": "success",
        "your_site_analysis": your_site_data,
        "competitor_analyses": competitor_data,
        "recommendations": recommendations["recommendations"],
        "priority_actions": recommendations["priority_actions"]
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "service": "Website Competitor Scanner"}


class ExportRequest(BaseModel):
    """Request model for docx export."""
    your_site_analysis: dict
    competitor_analyses: list[dict]
    recommendations: list[dict]
    priority_actions: list[dict]
    metric_insights: Optional[list[dict]] = None


@app.post("/api/export-docx")
async def export_docx(data: ExportRequest):
    """
    Generate a formatted Word document report that opens in Google Docs.
    """
    doc = Document()

    # Define styles
    title_style = doc.styles['Title']
    title_style.font.color.rgb = RGBColor(0x70, 0x82, 0x38)  # Olive green

    heading1_style = doc.styles['Heading 1']
    heading1_style.font.color.rgb = RGBColor(0x70, 0x82, 0x38)

    heading2_style = doc.styles['Heading 2']
    heading2_style.font.color.rgb = RGBColor(0xec, 0x58, 0x00)  # Persimmon

    # Title
    doc.add_heading('Website Optimization Report', 0)

    # Date and URL
    analysis = data.your_site_analysis
    seo = analysis.get("seo_factors", {})

    date_para = doc.add_paragraph()
    date_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}").italic = True
    date_para.add_run(f"\nURL Analyzed: {analysis.get('url', 'N/A')}")

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        "This report analyzes your website for SEO, GEO (Generative Engine Optimization), "
        "and LLM discoverability factors. Below are prioritized recommendations to improve "
        "your visibility in both traditional search and AI-powered search engines."
    )

    # Priority Actions
    doc.add_heading('Top Priority Actions', level=1)

    for i, action in enumerate(data.priority_actions or [], 1):
        # Action title with number
        p = doc.add_paragraph()
        p.add_run(f"{i}. {action.get('title', 'Untitled')}").bold = True

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Category: ").bold = True
        meta.add_run(f"{action.get('category', 'N/A')} | ")
        meta.add_run(f"Impact: ").bold = True
        meta.add_run(f"{action.get('impact', 'N/A')} | ")
        meta.add_run(f"Effort: ").bold = True
        meta.add_run(f"{action.get('effort', 'N/A')}")

        # Description
        if action.get('description'):
            doc.add_paragraph(action['description'])

        # Action steps
        all_actions = action.get('all_actions', [])
        if not all_actions and action.get('first_step'):
            all_actions = [action['first_step']]

        if all_actions:
            doc.add_paragraph("Action Steps:").runs[0].bold = True
            for step in all_actions:
                doc.add_paragraph(step, style='List Bullet')

        doc.add_paragraph()  # Spacing

    # Site Analysis Metrics
    doc.add_heading('Site Analysis Metrics', level=1)

    # SEO Factors table
    doc.add_heading('SEO Factors', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    seo_rows = [
        ('Title', seo.get('title', 'Missing')[:60] + '...' if len(seo.get('title', '')) > 60 else seo.get('title', 'Missing')),
        ('Title Length', f"{seo.get('title_length', 0)} chars (optimal: 50-60)"),
        ('Meta Description Length', f"{seo.get('meta_description_length', 0)} chars (optimal: 150-160)"),
        ('H1 Tags', f"{len(seo.get('h1_tags', []))} (optimal: 1)"),
        ('Word Count', str(seo.get('word_count', 0))),
        ('Images Missing Alt Text', str(seo.get('images_without_alt', 0)))
    ]

    for i, (label, value) in enumerate(seo_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    # Technical Factors
    doc.add_heading('Technical Factors', level=2)
    tech = analysis.get("technical_factors", {})
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    tech_rows = [
        ('HTTPS', 'Yes' if tech.get('https') else 'No'),
        ('Sitemap', 'Yes' if tech.get('has_sitemap') else 'No'),
        ('Robots.txt', 'Yes' if tech.get('has_robots_txt') else 'No'),
        ('Mobile Viewport', 'Yes' if tech.get('mobile_friendly_hints') else 'No')
    ]

    for i, (label, value) in enumerate(tech_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # LLM Discoverability
    doc.add_heading('LLM Discoverability', level=2)
    llm = analysis.get("llm_discoverability", {})
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    llm_rows = [
        ('Structured Content', 'Yes' if llm.get('structured_content') else 'No'),
        ('FAQ Schema', 'Yes' if llm.get('faq_schema') else 'No'),
        ('How-To Schema', 'Yes' if llm.get('how_to_schema') else 'No'),
        ('External Citations', str(llm.get('citations_and_sources', 0)))
    ]

    for i, (label, value) in enumerate(llm_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # GEO Factors
    doc.add_heading('GEO (AI Citation) Factors', level=2)
    geo = analysis.get("geo_factors", {})
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'

    geo_rows = [
        ('Citation Ready', 'Yes' if geo.get('citation_ready') else 'No'),
        ('Statistics Present', 'Yes' if geo.get('statistics_present') else 'No'),
        ('Comparison Tables', 'Yes' if geo.get('comparison_tables') else 'No'),
        ('Lists/Bullet Points', str(geo.get('lists_and_bullets', 0)))
    ]

    for i, (label, value) in enumerate(geo_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value

    # Metric Insights (if available)
    if data.metric_insights:
        doc.add_heading('Key Insights vs Competitors', level=1)
        for insight in data.metric_insights:
            p = doc.add_paragraph()
            p.add_run(f"{insight.get('metric', 'Unknown').replace('_', ' ').title()}").bold = True
            p.add_run(f" - {insight.get('status', 'unknown').upper()}")

            doc.add_paragraph(f"Your value: {insight.get('your_value', 'N/A')} | Competitor average: {insight.get('competitor_avg', 'N/A')}")

            if insight.get('explanation'):
                exp_para = doc.add_paragraph()
                exp_para.add_run("Why it matters: ").italic = True
                exp_para.add_run(insight['explanation'])

            if insight.get('recommendation'):
                rec_para = doc.add_paragraph()
                rec_para.add_run("Recommendation: ").bold = True
                rec_para.add_run(insight['recommendation'])

            doc.add_paragraph()

    # All Recommendations
    doc.add_heading('All Recommendations', level=1)

    for i, rec in enumerate(data.recommendations or [], 1):
        # Recommendation title
        p = doc.add_paragraph()
        p.add_run(f"{i}. [{rec.get('category', 'General')}] {rec.get('title', 'Untitled')}").bold = True

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Impact: {rec.get('impact', 'N/A')} | Effort: {rec.get('effort', 'N/A')}")

        # Description
        if rec.get('description'):
            doc.add_paragraph(rec['description'])

        # Action steps
        if rec.get('specific_actions'):
            doc.add_paragraph("Action Steps:").runs[0].bold = True
            for step in rec['specific_actions']:
                doc.add_paragraph(step, style='List Bullet')

        # Expected outcome
        if rec.get('expected_outcome'):
            outcome_para = doc.add_paragraph()
            outcome_para.add_run("Expected Outcome: ").bold = True
            outcome_para.add_run(rec['expected_outcome'])

        doc.add_paragraph()  # Spacing

    # Issues Found
    issues = analysis.get('issues', [])
    if issues:
        doc.add_heading('Issues Found', level=1)
        for issue in issues:
            p = doc.add_paragraph(style='List Bullet')
            severity = issue.get('severity', 'medium').upper()
            p.add_run(f"[{severity}] ").bold = True
            p.add_run(f"{issue.get('issue', '')} ({issue.get('category', 'General')})")

    # Strengths
    strengths = analysis.get('strengths', [])
    if strengths:
        doc.add_heading('Strengths', level=1)
        for strength in strengths:
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(f"{strength.get('strength', '')} ({strength.get('category', 'General')})")

    # Competitor Comparison
    if data.competitor_analyses:
        doc.add_heading('Competitor Comparison', level=1)

        # Create comparison table
        num_comps = len(data.competitor_analyses)
        table = doc.add_table(rows=7, cols=2 + num_comps)
        table.style = 'Table Grid'

        # Header row
        table.rows[0].cells[0].text = 'Metric'
        table.rows[0].cells[1].text = 'Your Site'
        for i, comp in enumerate(data.competitor_analyses):
            table.rows[0].cells[2 + i].text = comp.get('domain', comp.get('url', f'Competitor {i+1}'))[:20]

        # Data rows
        metrics = [
            ('Word Count', lambda a: str(a.get('seo_factors', {}).get('word_count', 'N/A'))),
            ('Structured Data', lambda a: '✓' if a.get('content_analysis', {}).get('has_structured_data') else '✗'),
            ('FAQ Schema', lambda a: '✓' if a.get('llm_discoverability', {}).get('faq_schema') else '✗'),
            ('Statistics Present', lambda a: '✓' if a.get('geo_factors', {}).get('statistics_present') else '✗'),
            ('Citation Ready', lambda a: '✓' if a.get('geo_factors', {}).get('citation_ready') else '✗'),
            ('Issues Found', lambda a: str(len(a.get('issues', []))))
        ]

        for row_idx, (metric_name, get_value) in enumerate(metrics, 1):
            table.rows[row_idx].cells[0].text = metric_name
            table.rows[row_idx].cells[1].text = get_value(analysis)
            for comp_idx, comp in enumerate(data.competitor_analyses):
                table.rows[row_idx].cells[2 + comp_idx].text = get_value(comp)

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Report generated by Website Competitor Scanner").italic = True

    # Save to bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Generate filename
    filename = f"website-optimization-report-{datetime.now().strftime('%Y-%m-%d')}.docx"

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
