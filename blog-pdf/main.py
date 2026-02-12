"""
Blog-to-PDF API
FastAPI app: accepts blog URL + brand doc, returns structured preview + base64 PDF.
"""

import base64
import os
import re
import tempfile
from typing import Optional

from fastapi import FastAPI, Form, File, UploadFile, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from dotenv import load_dotenv
from typing import List

load_dotenv()

# Import our modules
from scraper import fetch_blog
from extractor import extract_brief
from pdf_generator import generate_pdf

# For document parsing (reuse web-scanner logic)
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

app = FastAPI(
    title="Blog to PDF",
    description="Convert blog posts into Replica-branded executive briefs",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Serve static files
STATIC_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


@app.get("/")
async def home():
    return FileResponse(os.path.join(STATIC_DIR, "index.html"))


@app.get("/api/health")
async def health():
    return {"status": "healthy", "service": "blog-to-pdf"}


@app.post("/api/generate")
async def generate(
    blog_url: str = Form(...),
    page_preference: str = Form("2"),
    brand_docs: List[UploadFile] = File(default=[]),
):
    """
    Main endpoint. Accepts blog URL + optional brand docs (multiple allowed).
    Returns JSON: { status, pdf_b64, filename, extracted }
    """
    # ── 1. Fetch blog ─────────────────────────────────────────────────────────
    blog_data = await fetch_blog(blog_url)
    if blog_data.get("status") == "error":
        raise HTTPException(
            status_code=422,
            detail=f"Could not fetch blog: {blog_data.get('error', 'Unknown error')}"
        )

    blog_text  = blog_data.get("text", "")
    blog_title = blog_data.get("title", "")
    inline_images = blog_data.get("inline_images", [])

    if len(blog_text.split()) < 50:
        raise HTTPException(
            status_code=422,
            detail="Could not extract enough text from the blog URL. "
                   "The page may require JavaScript or block scrapers."
        )

    # ── 2. Extract brand docs text (combine all uploaded docs) ────────────────
    brand_doc_text = ""
    brand_doc_paths = []   # keep (tmp_path, filename) alive for verbatim extraction
    valid_docs = [d for d in brand_docs if d and d.filename]
    for brand_doc in valid_docs:
        tmp = tempfile.NamedTemporaryFile(
            suffix=os.path.splitext(brand_doc.filename)[1],
            delete=False
        )
        content = await brand_doc.read()
        tmp.write(content)
        tmp.flush()
        tmp.close()
        brand_doc_paths.append((tmp.name, brand_doc.filename))
        doc_text = _extract_doc_text(tmp.name, brand_doc.filename)
        if doc_text:
            brand_doc_text += f"\n\n--- {brand_doc.filename} ---\n{doc_text}"

    # ── 3. LLM extraction ─────────────────────────────────────────────────────
    pages = int(page_preference) if page_preference in ("2", "3") else 2

    try:
        extracted = await extract_brief(
            blog_text=blog_text,
            blog_title=blog_title,
            brand_doc_text=brand_doc_text,
            page_preference=pages,
            inline_images=inline_images
        )
    except Exception as e:
        # Clean up temp files before raising
        for tmp_path, _ in brand_doc_paths:
            try: os.unlink(tmp_path)
            except Exception: pass
        raise HTTPException(status_code=500, detail=f"AI extraction failed: {str(e)}")

    # ── 4. Override elevator pitch + CTA with verbatim brand doc content ─────
    # Try DOCX-native extraction first (uses paragraph styles for clean boundaries),
    # then fall back to flat-text regex for PDFs / TXT files.
    verbatim = {}
    for tmp_path, orig_filename in brand_doc_paths:
        ext = os.path.splitext(orig_filename.lower())[1]
        if ext in {".docx", ".doc"}:
            verbatim = _extract_brand_verbatim_docx(tmp_path) or {}
        if verbatim.get("elevator_pitch_body"):
            break   # found it — stop looking
    if not verbatim.get("elevator_pitch_body"):
        verbatim = _extract_brand_verbatim(brand_doc_text)

    # Clean up temp files now that we're done with them
    for tmp_path, _ in brand_doc_paths:
        try: os.unlink(tmp_path)
        except Exception: pass

    if verbatim.get("elevator_pitch_body"):
        extracted["elevator_pitch_body"] = verbatim["elevator_pitch_body"]
    if verbatim.get("elevator_pitch_header"):
        extracted["elevator_pitch_header"] = verbatim["elevator_pitch_header"]
    if verbatim.get("cta_text"):
        extracted["cta_text"] = verbatim["cta_text"]
    if verbatim.get("cta_url"):
        extracted["cta_url"] = verbatim["cta_url"]

    # ── 4b. Blog CTA fallback — use blog's own CTA link if brand doc didn't supply one
    blog_cta_text = blog_data.get("cta_text", "")
    blog_cta_url  = blog_data.get("cta_url", "")
    if blog_cta_url and not extracted.get("cta_url"):
        extracted["cta_url"]  = blog_cta_url
        if not extracted.get("cta_text"):
            extracted["cta_text"] = blog_cta_text

    # ── 5. Generate PDF ────────────────────────────────────────────────────────
    extracted["blog_url"] = blog_data.get("url", "")
    try:
        pdf_bytes = generate_pdf(extracted)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    # ── 6. Build filename ──────────────────────────────────────────────────────
    slug = re.sub(r"[^a-z0-9]+", "-", extracted.get("title", "brief").lower()).strip("-")[:40]
    filename = f"replica-brief-{slug}.pdf"

    pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")

    return {
        "status": "success",
        "pdf_b64": pdf_b64,
        "filename": filename,
        "extracted": extracted,
        "blog_meta": {
            "url": blog_data.get("url"),
            "title": blog_title,
            "word_count": blog_data.get("word_count", 0),
            "inline_images": inline_images
        }
    }


@app.post("/api/regenerate-with-image")
async def regenerate_with_image(
    brief_json: str = Form(...),
    section_index: int = Form(0),
    image_file: Optional[UploadFile] = File(default=None),
):
    """
    Regenerate PDF with a user-provided image embedded at a specific section.
    """
    import json

    try:
        data = json.loads(brief_json)
    except Exception:
        raise HTTPException(status_code=422, detail="Invalid brief JSON")

    image_paths = {}
    if image_file and image_file.filename:
        ext = os.path.splitext(image_file.filename)[1].lower()
        if ext not in {".png", ".jpg", ".jpeg", ".gif", ".webp"}:
            raise HTTPException(status_code=422, detail="Unsupported image format")

        tmp = tempfile.NamedTemporaryFile(suffix=ext, delete=False)
        try:
            content = await image_file.read()
            tmp.write(content)
            tmp.flush()
            tmp.close()
            image_paths[section_index] = tmp.name

            pdf_bytes = generate_pdf(data, image_paths=image_paths)
        finally:
            try:
                os.unlink(tmp.name)
            except Exception:
                pass
    else:
        pdf_bytes = generate_pdf(data)

    slug = re.sub(r"[^a-z0-9]+", "-", data.get("title", "brief").lower()).strip("-")[:40]
    filename = f"replica-brief-{slug}.pdf"

    return {
        "status": "success",
        "pdf_b64": base64.b64encode(pdf_bytes).decode("utf-8"),
        "filename": filename
    }


# ── Helpers ───────────────────────────────────────────────────────────────────

def _extract_brand_verbatim_docx(file_path: str) -> dict:
    """
    Extract elevator pitch and CTA from a DOCX file using paragraph styles.
    Looks for a heading whose text matches 'Boilerplate', 'Elevator pitch', etc.,
    then collects all following Normal/body paragraphs until the next Heading.
    This is far more reliable than regex on flattened text because it uses the
    document's own structure to define section boundaries.
    """
    result = {}
    try:
        import docx as docx_lib
        doc = docx_lib.Document(file_path)
    except Exception:
        return result

    heading_styles = {"heading 1", "heading 2", "heading 3", "heading 4"}
    pitch_labels = {
        "boilerplate", "elevator pitch", "short pitch", "about replica",
        "why replica", "our pitch", "one-liner", "value proposition",
        "company description", "positioning statement"
    }
    url_pattern = re.compile(r"https?://[^\s\)\]\"'<>]+")
    cta_labels = re.compile(
        r"^(cta|call[- ]to[- ]action|request\s+a\s+demo|book\s+a\s+demo|get\s+started)",
        re.IGNORECASE
    )

    paragraphs = doc.paragraphs
    i = 0
    while i < len(paragraphs):
        p = paragraphs[i]
        style_name = p.style.name.lower()
        text = p.text.strip()

        # Check if this is a heading that matches a pitch label
        if style_name in heading_styles and text.lower() in pitch_labels:
            # Collect all following paragraphs until the next heading
            body_parts = []
            i += 1
            while i < len(paragraphs):
                np = paragraphs[i]
                np_style = np.style.name.lower()
                np_text = np.text.strip()
                # Stop at any heading
                if np_style in heading_styles:
                    break
                if np_text:
                    body_parts.append(np_text)
                i += 1
            if body_parts:
                result["elevator_pitch_body"] = " ".join(body_parts).strip()
            break
        i += 1

    # CTA: scan entire doc for a URL near demo/contact language
    for p in paragraphs:
        text = p.text.strip()
        if not text:
            continue
        url_m = url_pattern.search(text)
        if url_m:
            url = url_m.group(0).rstrip(".,;)")
            if re.search(r"(demo|contact|get[-_]started|book|schedule|request)", url, re.I):
                if not result.get("cta_url"):
                    result["cta_url"] = url
                    # Text before the URL on the same line is the CTA label
                    label = text[:url_m.start()].strip(" -:")
                    if label and not result.get("cta_text"):
                        result["cta_text"] = label
        # Also check for explicit CTA label lines without URLs
        if cta_labels.match(text) and not result.get("cta_text"):
            remainder = re.sub(r"^[^:]+:\s*", "", text).strip()
            if remainder and not url_pattern.search(remainder):
                result["cta_text"] = remainder[:80]

    return result


def _extract_brand_verbatim(brand_doc_text: str) -> dict:
    """
    Extract elevator pitch and CTA from brand doc text verbatim (no LLM).
    Looks for common patterns: sections labelled "elevator pitch", "about us",
    "why replica", "our pitch", and CTA links/buttons.

    Returns dict with keys: elevator_pitch_header, elevator_pitch_body, cta_text, cta_url
    (any may be empty string if not found).
    """
    if not brand_doc_text:
        return {}

    result = {}

    # ── Elevator pitch / Boilerplate ───────────────────────────────────────────
    # Pass 1: look for "Boilerplate" label specifically — highest priority.
    boilerplate_pattern = re.compile(
        r"(?m)^[#*_\s]*boilerplate[#*_\s:]*$",
        re.IGNORECASE
    )
    match = boilerplate_pattern.search(brand_doc_text)

    # Pass 2: fallback to other pitch-like labels if no Boilerplate section found.
    if not match:
        fallback_pattern = re.compile(
            r"(?m)^[#*_\s]*"
            r"(elevator\s+pitch|why\s+replica|about\s+replica|our\s+pitch|"
            r"one[- ]liner|value\s+prop(?:osition)?|company\s+description|"
            r"positioning\s+statement|messaging\s+framework)"
            r"[#*_\s:]*$",
            re.IGNORECASE
        )
        match = fallback_pattern.search(brand_doc_text)

    if match:
        # Grab text following the header.
        # Allow up to 1 consecutive blank line (multi-sentence paragraphs);
        # stop at a second consecutive blank or a new section header.
        after = brand_doc_text[match.end():]
        lines = after.split("\n")
        body_lines = []
        blank_streak = 0
        for line in lines:
            stripped = line.strip()
            if not stripped:
                blank_streak += 1
                if blank_streak > 1:   # two blanks in a row = section boundary
                    break
                continue
            blank_streak = 0
            # Stop at next section header (# heading or ALL CAPS line)
            if re.match(r"^(#{1,3}|[A-Z][A-Z\s]{5,}$)", stripped):
                break
            body_lines.append(stripped)
        if body_lines:
            result["elevator_pitch_body"] = " ".join(body_lines).strip()

    # ── CTA text + URL ─────────────────────────────────────────────────────────
    # Look for CTA labels like "CTA:", "Call to action:", "Button:", "Link:"
    cta_label_pattern = re.compile(
        r"(?i)\b(cta|call[- ]to[- ]action|button|primary\s+cta|demo\s+link|"
        r"request\s+a\s+demo|book\s+a\s+demo|schedule\s+a\s+demo|get\s+started)"
        r"\s*[:\-]?\s*(.+)",
        re.IGNORECASE
    )
    url_pattern = re.compile(r"https?://[^\s\)\]\"'<>]+")

    for line in brand_doc_text.split("\n"):
        m = cta_label_pattern.search(line)
        if m:
            remainder = m.group(2).strip()
            # If the remainder contains a URL, split text from URL
            url_m = url_pattern.search(remainder)
            if url_m:
                if not result.get("cta_url"):
                    result["cta_url"] = url_m.group(0).rstrip(".,;)")
                text_part = remainder[:url_m.start()].strip(" -:")
                if text_part and not result.get("cta_text"):
                    result["cta_text"] = text_part
            elif remainder and not result.get("cta_text"):
                result["cta_text"] = remainder[:80]

    # Fallback: look for any URL that suggests a demo/contact page
    if not result.get("cta_url"):
        for url_m in url_pattern.finditer(brand_doc_text):
            url = url_m.group(0).rstrip(".,;)")
            if re.search(r"(demo|contact|get[-_]started|book|schedule|request)", url, re.I):
                result["cta_url"] = url
                break

    return result


def _extract_doc_text(file_path: str, original_filename: str) -> str:
    """Extract plain text from a document file (PDF, DOCX, TXT)."""
    ext = os.path.splitext(original_filename.lower())[1]

    try:
        if ext == ".pdf":
            import pypdf
            reader = pypdf.PdfReader(file_path)
            return "\n\n".join(
                p.extract_text() for p in reader.pages if p.extract_text()
            )
        elif ext in {".docx", ".doc"}:
            import docx
            doc = docx.Document(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
        else:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
    except Exception:
        return ""
